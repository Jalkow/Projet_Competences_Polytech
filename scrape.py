import os
import json
import re

from docx import Document
# UTILITAIRES

def clean_lines(text):
    return [l.strip() for l in text.split("\n") if l.strip()]


# META INFOS

def extract_title(doc):
    for p in doc.paragraphs:
        if p.text.strip():
            return p.text.strip()
    return None


def extract_specialite(doc):
    for p in doc.paragraphs:
        if "IDU" in p.text:
            return "IDU"
    return None


def extract_cours(doc):
    for p in doc.paragraphs:
        if "Proj" in p.text:
            return p.text.strip()
    return None


def extract_semestre(doc):
    for p in doc.paragraphs:
        match = re.search(r"S\d+", p.text)
        if match:
            return match.group()
    return None


# OBJECTIFS 

def extract_objectifs(doc):
    objectifs = []
    capture = False

    for p in doc.paragraphs:
        text = p.text.strip()

        if not text:
            continue

        
        if "l’objectif de ce projet est" in text.lower():
            capture = True

            
            if ":" in text:
                parts = text.split(":", 1)
                objectifs.append(parts[1].strip())
            else:
                objectifs.append(text)

            continue

        
        if capture:
            if "compétence" in text.lower():
                break
            objectifs.append(text)

    return objectifs



# COMPETENCES (TABLES)


def extract_competences(doc):
    competences = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header = " ".join([c.text.lower() for c in table.rows[0].cells])

        if "compétence" in header or "competence" in header:

            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]

                if len(cells) >= 2:
                    competences.append({
                        "nom": cells[0],
                        "apprentissages": clean_lines(cells[1])
                    })

    return competences



# FILTRAGE COMPETENCES METIER

def filter_competences(competences):
    filtres = []

    for c in competences:
        nom_lower = c["nom"].lower()

        if any(keyword in nom_lower for keyword in [
            "vérification",
            "réflexion",
            "compréhension",
            "argumenter",
            "entretien",
            "évaluation",
            "objectifs"
        ]):
            continue

        filtres.append(c)

    return filtres



# EVALUATION


def extract_evaluation(doc):
    evaluation = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header = " ".join([c.text.lower() for c in table.rows[0].cells])

        if "objectif" in header and "livrable" in header:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]

                if len(cells) >= 4:
                    evaluation.append({
                        "objectif": cells[0],
                        "livrable": cells[1],
                        "type": cells[2],
                        "outil": cells[3]
                    })

    return evaluation



# PARSER GLOBAL


def parse_docx(file_path):
    doc = Document(file_path)

    competences = extract_competences(doc)

    data = {
        "titre": extract_title(doc),
        "specialite": extract_specialite(doc),
        "cours": extract_cours(doc),
        "semestre": extract_semestre(doc),

        "objectifs": extract_objectifs(doc),

        "competences": filter_competences(competences),

        "evaluation": extract_evaluation(doc)
    }

    return data



# MAIN


def main():
    file_path = r"C:\Users\shogu\Downloads\scraper\APP IDU S6\20251029 AMS_APPS6_IDU fiche enseignant.docx"

    print( "Démarrage")

    if not os.path.exists(file_path):
        print(" Fichier introuvable :", file_path)
        return

    data = parse_docx(file_path)

    output_path = os.path.join(os.path.dirname(file_path), "output.json")

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

    print(" JSON généré :", output_path)


if __name__ == "__main__":
    main()